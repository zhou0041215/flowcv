from __future__ import annotations

from io import BytesIO
from pathlib import Path

from fastapi import UploadFile

from app.core.exceptions import AppException

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp", ".bmp"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r", "\n").splitlines()]
    return "\n".join(line for line in lines if line)


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise AppException("当前环境未安装 PDF 解析组件，请安装 pypdf 后重试", 503) from exc

    try:
        reader = PdfReader(BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        raise AppException("PDF 简历解析失败，请确认文件未损坏", 400) from exc


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise AppException("当前环境未安装 Word 解析组件，请安装 python-docx 后重试", 503) from exc

    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise AppException("Word 简历解析失败，请确认文件未损坏或改用 PDF 导入", 400) from exc
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_image(data: bytes) -> str:
    try:
        from PIL import Image, ImageFilter, ImageOps
        import pytesseract
    except ModuleNotFoundError as exc:
        raise AppException("当前环境未配置 OCR 组件，暂时无法识别图片简历，请安装 Pillow、pytesseract 和 OCR 语言包后重试", 503) from exc

    try:
        image = Image.open(BytesIO(data))
        image.load()
    except Exception as exc:
        raise AppException("图片文件读取失败，请确认文件未损坏或改用 PDF/Word 导入", 400) from exc

    image = image.convert("RGB")
    max_side = max(image.size)
    if max_side and max_side < 1800:
        scale = min(3, 1800 / max_side)
        resampling = getattr(getattr(Image, "Resampling", Image), "LANCZOS")
        image = image.resize((int(image.width * scale), int(image.height * scale)), resampling)

    processed = ImageOps.grayscale(image)
    processed = ImageOps.autocontrast(processed)
    processed = processed.filter(ImageFilter.SHARPEN)

    errors: list[str] = []
    for lang in ("chi_sim+eng", "chi_sim", "eng", ""):
        try:
            kwargs = {"config": "--psm 6"}
            if lang:
                kwargs["lang"] = lang
            text = pytesseract.image_to_string(processed, **kwargs)
            if text and text.strip():
                return text
        except Exception as exc:
            errors.append(str(exc))

    detail = "；".join(error for error in errors if error)[:180]
    if "tesseract is not installed" in detail or "not in your PATH" in detail:
        raise AppException("服务器未安装 Tesseract OCR，暂时无法导入图片简历。请安装 tesseract 与中文语言包后重试，或改用 PDF/Word 导入", 503)
    message = "图片简历 OCR 识别失败，请确认图片清晰、文字方向正确，或改用 PDF/Word 导入"
    if detail:
        message = f"{message}（OCR 详情：{detail}）"
    raise AppException(message, 400)


def _file_suffix(file: UploadFile) -> str:
    return Path(file.filename or "resume").suffix.lower()


def extract_resume_text_from_upload(file: UploadFile) -> str:
    filename = file.filename or "resume"
    suffix = _file_suffix(file)
    if suffix not in SUPPORTED_EXTENSIONS:
        raise AppException("仅支持 PDF、Word(.docx)、图片、TXT/Markdown 简历导入", 400)

    data = file.file.read()
    if not data:
        raise AppException("上传文件为空", 400)

    if suffix == ".pdf":
        text = _extract_pdf(data)
    elif suffix == ".docx":
        text = _extract_docx(data)
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", "ignore")
    else:
        text = _extract_image(data)

    text = _clean_text(text)
    if len(text) < 20:
        raise AppException("没有识别到可导入的简历内容，请确认文件清晰且包含文本", 400)
    return text[:30000]


def extract_resume_text_from_uploads(files: list[UploadFile]) -> str:
    actual_files = [file for file in files if file and file.filename]
    if not actual_files:
        raise AppException("请上传要导入的简历文件", 400)

    if len(actual_files) == 1:
        return extract_resume_text_from_upload(actual_files[0])

    invalid_files = [file.filename or "未知文件" for file in actual_files if _file_suffix(file) not in IMAGE_EXTENSIONS]
    if invalid_files:
        raise AppException("多文件导入仅支持多张图片。PDF、Word、TXT/Markdown 请单独上传", 400)

    parts: list[str] = []
    for index, file in enumerate(actual_files, start=1):
        data = file.file.read()
        if not data:
            continue
        text = _clean_text(_extract_image(data))
        if text:
            parts.append(f"第 {index} 张图片（{file.filename or '未命名图片'}）识别内容：\n{text}")

    merged_text = _clean_text("\n\n".join(parts))
    if len(merged_text) < 20:
        raise AppException("没有识别到可导入的简历内容，请确认图片清晰、文字方向正确，或改用 PDF/Word 导入", 400)
    return merged_text[:30000]
