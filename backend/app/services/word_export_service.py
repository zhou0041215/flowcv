from pathlib import Path
from typing import Any

from docx import Document
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.export_record import ExportRecord
from app.models.resume import Resume
from app.services.resume_locale import localized_section_title, resolve_resume_language, resume_locale, section_field_labels
from app.services.rich_text_service import rich_text_to_plain


def _plain(value: Any, colon: str = "：") -> str:
    if isinstance(value, list):
        lines = []
        for item in value:
            if isinstance(item, dict):
                text = " / ".join(_inline(part) for part in item.values() if _inline(part))
            else:
                text = rich_text_to_plain(item)
            if text:
                lines.append(text)
        return "\n".join(lines)
    if isinstance(value, dict):
        return "\n".join(f"{key}{colon}{_inline(item)}" for key, item in value.items() if _inline(item))
    return rich_text_to_plain(value)


def _inline(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return " / ".join(_inline(item) for item in value if _inline(item))
    if isinstance(value, dict):
        return " / ".join(_inline(item) for item in value.values() if _inline(item))
    return rich_text_to_plain(value)


def _inline_description(value: Any) -> str:
    lines = []
    for line in rich_text_to_plain(value).splitlines():
        text = line.lstrip(" \t•·-*").strip()
        if text:
            lines.append(text)
    return "；".join(lines)


def _has_text(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return bool(rich_text_to_plain(value).strip())
    if isinstance(value, list):
        return any(_has_text(item) for item in value)
    if isinstance(value, dict):
        return any(_has_text(item) for key, item in value.items() if key not in {"id", "preset_type"})
    return bool(value)


def export_word(db: Session, user_id: int, resume: Resume) -> Path:
    settings.export_path.mkdir(parents=True, exist_ok=True)
    data = resume.resume_data or {}
    layout = data.get("layout", {}) if isinstance(data.get("layout"), dict) else {}
    basics = data.get("basics", {})
    language = resolve_resume_language(resume.language, data)
    labels = resume_locale(language)
    colon = labels["colon"]
    doc = Document()
    doc.add_heading(_inline(basics.get("name")) or resume.title, 0)
    doc.add_paragraph(" | ".join(filter(None, [_inline(basics.get("title")), _inline(basics.get("phone")), _inline(basics.get("email")), _inline(basics.get("location"))])))

    titles = layout.get("section_titles", {})
    skills_options = layout.get("skills_options") if isinstance(layout.get("skills_options"), dict) else {}
    show_skill_keywords = skills_options.get("show_keywords") is not False
    skill_description_inline = skills_options.get("description_inline") is True
    custom_sections = {
        item["id"]: item
        for item in data.get("custom_sections", [])
        if isinstance(item, dict) and item.get("id")
    }
    for key in layout.get("section_order", []):
        if key == "basics" or key in layout.get("hidden_sections", []):
            continue
        section = custom_sections.get(key) if key in custom_sections else data.get(key)
        if not section or not _has_text(section):
            continue
        doc.add_heading(localized_section_title(key, titles.get(key), language), level=1)
        if key == "summary":
            doc.add_paragraph(_plain(section.get("content"), colon))
        elif key == "skills" and isinstance(section, list):
            field_labels = section_field_labels(layout, key, language)
            keywords_label = field_labels.get("keywords", "")
            for item in section:
                name = _inline(item.get("name"))
                description = _inline_description(item.get("description")) if skill_description_inline else _plain(item.get("description"), colon)
                if skill_description_inline and show_skill_keywords and item.get("keywords"):
                    prefix = f"{keywords_label}{colon}" if keywords_label else ""
                    doc.add_paragraph(f"{prefix}{_inline(item.get('keywords'))}")
                if skill_description_inline and description:
                    doc.add_paragraph(f"{name}{colon}{description}" if name else description)
                elif name:
                    doc.add_heading(name, level=2)
                if not skill_description_inline and show_skill_keywords and item.get("keywords"):
                    prefix = f"{keywords_label}{colon}" if keywords_label else ""
                    doc.add_paragraph(f"{prefix}{_inline(item.get('keywords'))}")
                if description and not skill_description_inline:
                    doc.add_paragraph(description)
        elif isinstance(section, list):
            field_labels = section_field_labels(layout, key, language)
            for item in section:
                doc.add_heading(
                    _inline(item.get("name") or item.get("company") or item.get("school") or item.get("title"))
                    or labels["experience"],
                    level=2,
                )
                meta = " / ".join(
                    filter(
                        None,
                        [
                            _inline(item.get("role")),
                            _inline(item.get("position")),
                            _inline(item.get("major")),
                            _inline(item.get("degree")),
                        ],
                    )
                )
                if meta:
                    doc.add_paragraph(meta)
                if item.get("tech_stack"):
                    tech_label = field_labels.get("tech_stack", "")
                    prefix = f"{tech_label}{colon}" if tech_label else ""
                    doc.add_paragraph(f"{prefix}{_inline(item.get('tech_stack'))}")
                for field in ["description", "highlights", "content"]:
                    if item.get(field):
                        doc.add_paragraph(_plain(item.get(field), colon))
        elif key.startswith("custom_"):
            field_labels = section_field_labels(layout, key, language, custom=True)
            for item in section.get("items", []):
                heading = _inline(
                    item.get("title")
                    or item.get("name")
                    or item.get("company")
                    or item.get("organization")
                    or item.get("publisher")
                    or item.get("platform")
                )
                doc.add_heading(heading or labels["content"], level=2)
                meta = " / ".join(
                    filter(
                        None,
                        [
                            _inline(item.get("issuer")),
                            _inline(item.get("level")),
                            _inline(item.get("score")),
                            _inline(item.get("award")),
                            _inline(item.get("role")),
                            _inline(item.get("position")),
                            _inline(item.get("institution")),
                            _inline(item.get("publisher")),
                            _inline(item.get("platform")),
                            _inline(item.get("date")),
                            _inline(item.get("start_date")),
                            _inline(item.get("end_date")),
                            _inline(item.get("credential_id")),
                            _inline(item.get("url")),
                        ],
                    )
                )
                if meta:
                    doc.add_paragraph(meta)
                if item.get("tech_stack"):
                    tech_label = field_labels.get("tech_stack", "")
                    prefix = f"{tech_label}{colon}" if tech_label else ""
                    doc.add_paragraph(f"{prefix}{_inline(item.get('tech_stack'))}")
                for field in ["content", "description", "highlights"]:
                    if item.get(field):
                        doc.add_paragraph(_plain(item.get(field), colon))

    file_name = f"flowcv_resume_{resume.id}.docx"
    file_path = settings.export_path / file_name
    doc.save(file_path)
    db.add(ExportRecord(user_id=user_id, resume_id=resume.id, file_type="word", file_name=file_name, file_path=str(file_path)))
    db.commit()
    return file_path
